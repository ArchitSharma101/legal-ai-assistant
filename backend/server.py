from fastapi import FastAPI, APIRouter, File, UploadFile, HTTPException, Form, Response
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import pymongo
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Any
from dataclasses import dataclass
import uuid
from datetime import datetime, timezone
import shutil
import httpx
import asyncio
from PyPDF2 import PdfReader

# Custom classes for message handling
@dataclass
class UserMessage:
    text: str
    file_contents: Optional[List[Any]] = None

@dataclass
class FileContentWithMimeType:
    file_path: str
    mime_type: str

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

logging.basicConfig(level=logging.INFO)
logging.info(f"GEMINI_API_KEY loaded: {'set' if os.environ.get('GEMINI_API_KEY') else 'NOT set'}")
if not os.environ.get('GEMINI_API_KEY'):
    logging.error("GEMINI_API_KEY is not set. Document analysis will not work!")

# Create the main app without a prefix
app = FastAPI()

# MongoDB connection
try:
    mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
    client = AsyncIOMotorClient(mongo_url)
    db = client[os.environ.get('DB_NAME', 'legal_docs')]
    # Test the connection
    client.admin.command('ping')
    logging.info("Successfully connected to MongoDB")
except Exception as e:
    logging.error(f"Failed to connect to MongoDB: {str(e)}")
    raise

# Add CORS middleware first
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],  # Allow all origins
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Create uploads directory
UPLOAD_DIR = ROOT_DIR / "../uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

# Pydantic Models
class Document(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    file_path: str
    file_type: str
    upload_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    analysis_status: str = "pending"  # pending, processing, completed, failed
    summary: Optional[str] = None
    key_clauses: Optional[List[dict]] = None
    risk_assessment: Optional[str] = None

class DocumentCreate(BaseModel):
    filename: str
    file_type: str

class QuestionRequest(BaseModel):
    document_id: str
    question: str
    session_id: Optional[str] = None

class ChatMessage(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str
    session_id: str
    question: str
    answer: str
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AnalysisRequest(BaseModel):
    document_id: str

# OCR helper function for PDFs
def extract_text_with_ocr(pdf_path: str) -> str:
    """Extract text from PDF using OCR when regular extraction fails"""
    try:
        # For now, return a message that OCR is not available
        return "This PDF appears to be image-based or scanned. OCR processing is not currently available in this environment. Please provide a text-based PDF or use external OCR services to convert the document to text format before uploading."
    except Exception as e:
        return f"OCR processing failed: {str(e)}"

GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')
if not GEMINI_API_KEY:
    logging.error("GEMINI_API_KEY environment variable is not set!")
    raise ValueError("GEMINI_API_KEY environment variable is required")

GEMINI_API_URL = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={GEMINI_API_KEY}"

async def send_message(prompt: str) -> str:
    if not GEMINI_API_KEY:
        raise HTTPException(status_code=503, detail="AI service is not configured. Please set GEMINI_API_KEY environment variable.")

    max_retries = 3

    for attempt in range(max_retries):
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    GEMINI_API_URL,
                    json={
                        "contents": [
                            {
                                "parts": [
                                    {
                                        "text": prompt
                                    }
                                ]
                            }
                        ]
                    },
                    headers={
                        "Content-Type": "application/json"
                    }
                )

                if response.status_code == 200:
                    data = response.json()
                    if 'candidates' in data and len(data['candidates']) > 0:
                        candidate = data['candidates'][0]
                        if 'content' in candidate and 'parts' in candidate['content'] and len(candidate['content']['parts']) > 0:
                            return candidate['content']['parts'][0]['text']
                        else:
                            raise Exception("Invalid response structure from Gemini API")
                    else:
                        raise Exception("No candidates in Gemini API response")
                elif response.status_code in (429, 500, 502, 503, 504):
                    if attempt < max_retries - 1:
                        wait_time = 2 ** attempt
                        logging.warning(f"Gemini API retryable error (attempt {attempt + 1}/{max_retries}): {response.status_code} - {response.text}")
                        await asyncio.sleep(wait_time)
                        continue
                    else:
                        logging.error(f"Gemini API failed after {max_retries} attempts: {response.status_code} - {response.text}")
                        raise HTTPException(status_code=503, detail=f"AI service is temporarily unavailable. Please try again later.")
                else:
                    logging.error(f"Gemini API non-retryable error: {response.status_code} - {response.text}")
                    raise HTTPException(status_code=500, detail=f"AI service error: {response.text}")

        except httpx.TimeoutException:
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt
                logging.warning(f"Gemini API timeout (attempt {attempt + 1}/{max_retries})")
                await asyncio.sleep(wait_time)
                continue
            else:
                logging.error(f"Gemini API timeout after {max_retries} attempts")
                raise HTTPException(status_code=504, detail="AI service request timed out. Please try again later.")

        except Exception as e:
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt
                logging.warning(f"Gemini API error (attempt {attempt + 1}/{max_retries}): {str(e)}")
                await asyncio.sleep(wait_time)
                continue
            else:
                logging.error(f"Gemini API failed after {max_retries} attempts: {str(e)}")
                raise HTTPException(status_code=500, detail=f"AI service error: {str(e)}")

    # This should never be reached, but just in case
    raise HTTPException(status_code=500, detail="AI service unavailable after retries")

@api_router.get("/")
async def root():
    return {"message": "Legal Document AI Assistant API"}

# Document upload endpoint
@api_router.post("/documents/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload a legal document for analysis"""
    logging.info(f"Received upload request for file: {file.filename}")
    try:
        # Validate file type
        allowed_types = ['application/pdf', 'text/plain', 'application/msword', 
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document']
        if file.content_type not in allowed_types:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type}. Allowed types: {', '.join(allowed_types)}")
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        file_extension = Path(file.filename).suffix
        unique_filename = f"{file_id}{file_extension}"
        file_path = UPLOAD_DIR / unique_filename
        
        # Read file content fully before writing
        contents = await file.read()
        
        # Save file
        try:
            logging.info(f"Attempting to save file to: {file_path}")
            with open(file_path, "wb") as buffer:
                buffer.write(contents)
            logging.info("File saved successfully")

            # Create document record
            document = Document(
                filename=file.filename,
                file_path=str(file_path),
                file_type=file.content_type
            )
            
            # Save to database (using synchronous approach to avoid event loop conflicts)
            document_dict = document.model_dump()
            document_dict['upload_date'] = document_dict['upload_date'].isoformat()

            # Use synchronous MongoDB client for this operation
            import pymongo
            sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
            sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
            sync_db.documents.insert_one(document_dict)
            sync_client.close()
            
            return {"document_id": document.id, "filename": document.filename, "status": "uploaded"}
            
        except Exception as e:
            logging.error(f"Error saving file: {str(e)}")
            # Clean up the file if it was created but there was a database error
            if file_path.exists():
                try:
                    file_path.unlink()
                except:
                    pass
            raise
        
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Upload error: {str(e)}")
        raise HTTPException(status_code=500, detail="File upload failed")

# Get all documents
@api_router.get("/documents", response_model=List[Document])
async def get_documents():
    """Get all uploaded documents"""
    try:
        # Use synchronous MongoDB client to avoid event loop conflicts
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        documents = list(sync_db.documents.find())
        sync_client.close()

        for doc in documents:
            if isinstance(doc.get('upload_date'), str):
                doc['upload_date'] = datetime.fromisoformat(doc['upload_date'])
        return [Document(**doc) for doc in documents]
    except Exception as e:
        logging.error(f"Get documents error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve documents")

# Delete document endpoint
@api_router.delete("/documents/{document_id}")
async def delete_document(document_id: str):
    """Delete a document by ID"""
    try:
        # Log the delete request
        logging.info(f"Attempting to delete document with ID: {document_id}")

        # Use synchronous MongoDB client to avoid event loop conflicts
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        document = sync_db.documents.find_one({"id": document_id})
        if not document:
            sync_client.close()
            logging.warning(f"Document not found for deletion: {document_id}")
            raise HTTPException(status_code=404, detail="Document not found")

        # Delete the file from the filesystem
        try:
            file_path = Path(document['file_path'])
            if file_path.exists():
                file_path.unlink()
                logging.info(f"File deleted successfully: {file_path}")
        except Exception as e:
            logging.error(f"Error deleting file: {e}")
            # Continue with document deletion even if file deletion fails

        # Delete the document record from the database
        result = sync_db.documents.delete_one({"id": document_id})
        sync_client.close()

        if result.deleted_count == 0:
            logging.warning(f"Document not deleted from database: {document_id}")
            raise HTTPException(status_code=404, detail="Document not found in database")

        logging.info(f"Document deleted successfully: {document_id}")
        return {"message": "Document deleted successfully", "id": document_id}

    except HTTPException as he:
        raise he
    except Exception as e:
        logging.error(f"Delete document error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Get specific document
@api_router.get("/documents/{document_id}")
async def get_document(document_id: str):
    """Get a specific document by ID"""
    try:
        # Use synchronous MongoDB client to avoid event loop conflicts
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        document = sync_db.documents.find_one({"id": document_id})
        sync_client.close()

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        if isinstance(document.get('upload_date'), str):
            document['upload_date'] = datetime.fromisoformat(document['upload_date'])

        return Document(**document)
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Get document error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve document")

# Analyze document
@api_router.post("/documents/{document_id}/analyze")
async def analyze_document(document_id: str):
    """Analyze a document using AI"""
    try:
        # Get document from database
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        document = sync_db.documents.find_one({"id": document_id})
        if not document:
            sync_client.close()
            raise HTTPException(status_code=404, detail="Document not found")

        # Check if analysis already completed
        if document.get('analysis_status') == 'completed' and document.get('summary'):
            sync_client.close()
            logging.info(f"Returning cached analysis for document {document_id}")
            return {
                "document_id": document_id,
                "status": "completed",
                "summary": document.get('summary', ''),
                "key_clauses": document.get('key_clauses', []),
                "risk_assessment": document.get('risk_assessment', '')
            }

        # Update status to processing
        sync_db.documents.update_one(
            {"id": document_id},
            {"$set": {"analysis_status": "processing"}}
        )
        sync_client.close()

        # Read the document content
        document_text = ""
        try:
            if document['file_type'] == 'application/pdf' or document['file_path'].lower().endswith('.pdf'):
                # Extract text from PDF using PyPDF2
                try:
                    reader = PdfReader(document['file_path'])
                    text_pages = []
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_pages.append(page_text)

                    document_text = "\n".join(text_pages)

                    if len(document_text.strip()) < 50:
                        # Try OCR as fallback for image-based PDFs
                        logging.info("Attempting OCR processing for image-based PDF")
                        ocr_text = extract_text_with_ocr(document['file_path'])

                        if len(ocr_text.strip()) > len(document_text.strip()):
                            document_text = ocr_text
                            logging.info(f"OCR extracted additional content. Total length: {len(document_text)} characters")
                        else:
                            document_text = f"This PDF document appears to be image-based or scanned. The text extraction yielded only {len(document_text)} characters, which is insufficient for meaningful analysis. OCR processing was attempted but did not yield better results. This document may contain images, forms, or scanned text that requires specialized OCR tools or manual transcription. For proper analysis, please provide a text-based PDF or use professional OCR services to convert the document to text format first."
                            logging.warning(f"PDF text extraction and OCR yielded insufficient content: {len(document_text)} characters")
                    else:
                        logging.info(f"Extracted text from PDF. Length: {len(document_text)} characters")
                        logging.info(f"PDF extracted text: {document_text[:500]}")

                except Exception as e:
                    document_text = f"Failed to extract text from PDF: {str(e)}. This PDF may be corrupted, password-protected, or in an unsupported format."
                    logging.error(f"PDF text extraction error for file {document['file_path']}: {str(e)}")
            else:
                with open(document['file_path'], 'r', encoding='utf-8') as f:
                    document_text = f.read()
                logging.info(f"Document content read successfully. Length: {len(document_text)} characters")
                logging.info(f"First 500 characters: {document_text[:500]}")

                # Validate that we have meaningful content
                if len(document_text.strip()) < 50:
                    document_text = f"Document appears to be too short or empty. Content length: {len(document_text)} characters. This may indicate a reading issue or an unsupported file format."
                    logging.warning(f"Document content is too short: {len(document_text)} characters")

        except UnicodeDecodeError:
            # If it's not text, try binary read and decode
            try:
                with open(document['file_path'], 'rb') as f:
                    document_bytes = f.read()
                    document_text = document_bytes.decode('utf-8')
                logging.info(f"Document content decoded from binary. Length: {len(document_text)} characters")
            except UnicodeDecodeError:
                document_text = f"Unable to decode document content. This appears to be a binary file that requires specialized text extraction. File type: {document['file_type']}"
                logging.error(f"Failed to decode document content for file: {document['file_path']}, type: {document['file_type']}")
        except FileNotFoundError:
            document_text = f"Document file not found at path: {document['file_path']}"
            logging.error(f"Document file not found: {document['file_path']}")
        except Exception as e:
            document_text = f"Error reading document: {str(e)}"
            logging.error(f"Error reading document content: {str(e)}")

        # Create file content for analysis
        file_content = FileContentWithMimeType(
            file_path=document['file_path'],
            mime_type=document['file_type']
        )

        # Prepare analysis prompt with document content
        if len(document_text.strip()) < 200 or "Error reading document" in document_text or "Failed to extract text" in document_text or "Document appears to be too short" in document_text:
            # If document content is too short, contains error messages, or is not meaningful legal content
            analysis_prompt = """I need you to provide a legal document analysis, but the document content could not be properly extracted or is insufficient for detailed analysis. This may be due to file format limitations, reading errors, or the document being too short."""
        else:
            # If we have meaningful document content, analyze it specifically
            analysis_prompt = f"""You are a senior legal document analysis expert with 20+ years of experience reviewing contracts and legal agreements. Below is the content of a legal document that requires thorough professional analysis.

DOCUMENT CONTENT TO ANALYZE:
{document_text[:12000]}

CRITICAL REQUIREMENTS:
- Base your ENTIRE analysis on the specific content provided above
- Do NOT provide generic legal advice or templates
- Reference specific text, terms, and clauses from THIS document
- Be precise, professional, and actionable in your analysis

ANALYSIS REQUEST:
Provide a comprehensive structured analysis with these exact section headers:

EXECUTIVE SUMMARY
Provide a detailed executive summary that includes:
- Document Type & Purpose: What kind of legal document this is and its primary objective
- Key Parties Involved: Who are the main parties and their roles
- Core Obligations: The fundamental commitments and responsibilities outlined
- Critical Timeline: Any important dates, deadlines, or time-sensitive provisions
- Financial Implications: Any monetary amounts, payment terms, or financial commitments mentioned
- Overall Risk Level: High-level assessment of potential risks (Low/Medium/High)

KEY CLAUSES ANALYSIS
Identify and analyze the 7 most critical clauses or provisions from THIS document's content. For each clause, provide:

1. Exact Clause Reference: Quote the specific section/clause number or title from the document
2. Plain English Explanation: What this clause means in simple, understandable terms
3. Legal Implications: What rights, obligations, or restrictions it creates
4. Business Impact: How this affects the parties' relationship or operations
5. Potential Concerns: Any ambiguities, one-sided terms, or areas requiring attention

Format each clause as a numbered item with clear subheadings.

RISK ASSESSMENT
Conduct a thorough risk analysis covering:

HIGH-RISK ISSUES
- Terms that could cause significant financial loss or legal liability
- Unfavorable or one-sided provisions that disadvantage one party
- Ambiguous language that could lead to disputes
- Missing protections or standard clauses that should be present

MEDIUM-RISK CONCERNS
- Terms that could create operational difficulties
- Provisions requiring ongoing compliance or monitoring
- Clauses with conditional obligations that depend on external factors

LOW-RISK ITEMS
- Standard boilerplate provisions
- Clearly defined, balanced terms
- Protective clauses that benefit both parties

RECOMMENDATIONS
- Specific actions to mitigate identified risks
- Suggested modifications or negotiations points
- Professional review recommendations

PLAIN ENGLISH EXPLANATION
Transform the complex legal language into clear, accessible explanations:

What This Document Does: Simple overview of the agreement's purpose and effect
What Each Party Must Do: Clear list of obligations for each party involved
What Happens If Things Go Wrong: Consequences of breach or default
Important Dates & Deadlines: Timeline of key events and requirements
Money Matters: All financial terms explained in plain language
Key Rights & Protections: What each party can and cannot do
How to Get Out: Termination, cancellation, or exit provisions

FORMATTING REQUIREMENTS:
- Use clear, professional formatting with proper spacing and structure
- Use markdown formatting for better readability (bold, italics, lists, etc.)
- Ensure consistent spacing between sections and subsections
- Make the analysis comprehensive yet accessible to non-lawyers
- Format lists and bullet points professionally
- Use clear headings and subheadings for easy navigation
- Do not use emojis or special characters in headings - use plain text only
- Ensure every conclusion is directly supported by the document content
- Provide well-phrased, clear explanations with good grammar and structure"""

        try:
            logging.info("Starting AI analysis call")
            # Get AI analysis
            analysis_text = await send_message(analysis_prompt)
            logging.info(f"AI Analysis response received. Length: {len(analysis_text)}")

            # Validate the response
            if not analysis_text or len(analysis_text.strip()) < 100:
                raise Exception("Analysis response too short or empty")

            import re
            # Parse sections using regex for more robust extraction

            # Extract executive summary
            summary = ""
            try:
                summary_match = re.search(r"EXECUTIVE SUMMARY(.*?)(?=KEY CLAUSES ANALYSIS|$)", analysis_text, re.DOTALL)
                if summary_match:
                    summary = summary_match.group(1).strip()
                else:
                    logging.warning("EXECUTIVE SUMMARY section not found in analysis text")
            except Exception as e:
                logging.error(f"Error parsing executive summary section with regex: {str(e)}")
                summary = ""

            # Extract and parse key clauses
            key_clauses = []
            try:
                key_clauses_match = re.search(r"KEY CLAUSES ANALYSIS(.*?)(?=RISK ASSESSMENT|$)", analysis_text, re.DOTALL)
                if key_clauses_match:
                    clauses_section = key_clauses_match.group(1).strip()
                    # Parse individual clauses (numbered 1-7)
                    clause_pattern = r"(\d+)\.\s*(.*?)(?=\d+\.|$)"
                    matches = re.findall(clause_pattern, clauses_section, re.DOTALL)
                    for match in matches[:7]:  # Limit to 7 clauses
                        clause_num, clause_content = match
                        # Parse sub-components within each clause
                        sub_parts = re.split(r'\n\s*(?:[A-Z][a-zA-Z\s]+:)', clause_content)
                        if len(sub_parts) >= 2:
                            key_clauses.append({
                                "clause_number": int(clause_num),
                                "reference": sub_parts[1].strip() if len(sub_parts) > 1 else "",
                                "explanation": sub_parts[2].strip() if len(sub_parts) > 2 else "",
                                "implications": sub_parts[3].strip() if len(sub_parts) > 3 else "",
                                "impact": sub_parts[4].strip() if len(sub_parts) > 4 else "",
                                "concerns": sub_parts[5].strip() if len(sub_parts) > 5 else ""
                            })
                else:
                    logging.warning("KEY CLAUSES ANALYSIS section not found in analysis text")
            except Exception as e:
                logging.error(f"Error parsing key clauses section with regex: {str(e)}")
                key_clauses = []

            risk_assessment = ""
            try:
                risk_match = re.search(r"RISK ASSESSMENT(.*?)(?=PLAIN ENGLISH EXPLANATION|$)", analysis_text, re.DOTALL)
                if risk_match:
                    risk_assessment = risk_match.group(1).strip()
                else:
                    logging.warning("RISK ASSESSMENT section not found in analysis text")
            except Exception as e:
                logging.error(f"Error parsing risk assessment section with regex: {str(e)}")
                risk_assessment = ""

        except Exception as e:
            logging.error(f"Error during AI analysis: {str(e)}", exc_info=True)
            await db.documents.update_one(
                {"id": document_id},
                {"$set": {"analysis_status": "failed"}}
            )
            raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

        # Update document with analysis
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        sync_db.documents.update_one(
            {"id": document_id},
            {"$set": {
                "analysis_status": "completed",
                "summary": summary,
                "key_clauses": key_clauses,
                "risk_assessment": risk_assessment.strip()
            }}
        )
        sync_client.close()

        return {
            "document_id": document_id,
            "status": "completed",
            "summary": summary,
            "key_clauses": key_clauses,
            "risk_assessment": risk_assessment.strip()
        }

    except Exception as e:
        logging.error(f"Analysis error: {e}")
        # Update status to failed
        await db.documents.update_one(
            {"id": document_id},
            {"$set": {"analysis_status": "failed"}}
        )
        raise HTTPException(status_code=500, detail=f"Document analysis failed: {str(e)}")

# Ask questions about document
@api_router.post("/documents/ask")
async def ask_question(request: QuestionRequest):
    """Ask a specific question about a document"""
    try:
        # Get document from database
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        document = sync_db.documents.find_one({"id": request.document_id})
        if not document:
            sync_client.close()
            raise HTTPException(status_code=404, detail="Document not found")
        sync_client.close()
        
        session_id = request.session_id or f"qa_{request.document_id}_{uuid.uuid4()}"

        # Create file content for analysis
        file_content = FileContentWithMimeType(
            file_path=document['file_path'],
            mime_type=document['file_type']
        )
        
        # Read the document content for the question
        try:
            if document['file_type'] == 'application/pdf' or document['file_path'].lower().endswith('.pdf'):
                # Extract text from PDF using PyPDF2
                try:
                    reader = PdfReader(document['file_path'])
                    text_pages = []
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_pages.append(page_text)

                    document_text = "\n".join(text_pages)

                    if len(document_text.strip()) < 50:
                        # Try OCR as fallback for image-based PDFs
                        logging.info("Attempting OCR processing for image-based PDF in question answering")
                        ocr_text = extract_text_with_ocr(document['file_path'])

                        if len(ocr_text.strip()) > len(document_text.strip()):
                            document_text = ocr_text
                            logging.info(f"OCR extracted additional content for question. Total length: {len(document_text)} characters")
                        else:
                            document_text = f"This PDF document appears to be image-based or scanned. The text extraction yielded only {len(document_text)} characters, which is insufficient for answering questions. OCR processing was attempted but did not yield better results. This document may contain images, forms, or scanned text that requires specialized OCR tools. For proper question answering, please provide a text-based PDF or use professional OCR services first."
                            logging.warning(f"PDF text extraction and OCR yielded insufficient content for question: {len(document_text)} characters")
                    else:
                        logging.info(f"Extracted text from PDF for question. Length: {len(document_text)} characters")

                except Exception as e:
                    document_text = f"Failed to extract text from PDF: {str(e)}. This PDF may be corrupted, password-protected, or in an unsupported format."
                    logging.error(f"PDF text extraction error for file {document['file_path']}: {str(e)}")
            else:
                with open(document['file_path'], 'r', encoding='utf-8') as f:
                    document_text = f.read()
                logging.info(f"Document content read for question. Length: {len(document_text)} characters")
        except UnicodeDecodeError:
            with open(document['file_path'], 'rb') as f:
                document_bytes = f.read()
                try:
                    document_text = document_bytes.decode('utf-8')
                except UnicodeDecodeError:
                    document_text = "Unable to decode document content. This may be a binary file that requires specialized processing."
        except Exception as e:
            document_text = f"Error reading document: {str(e)}"

        # Create question message with document content
        if len(document_text.strip()) < 100:
            # If document content is too short or empty
            question_prompt = f"""I need to ask a question about a legal document, but the document content could not be properly extracted. This may be due to file format limitations.

Document type: {document['file_type']}
Question: {request.question}

Since I cannot access the specific document content, please provide general guidance about this type of question for {document['file_type']} documents. Explain what factors would typically be considered when answering this question for this document type."""
        else:
            # If we have meaningful document content
            question_prompt = f"""You are a legal document analysis expert. Below is the content of a legal document. Please answer the specific question based ONLY on this document's content.

DOCUMENT CONTENT:
{document_text[:10000]}

QUESTION: {request.question}

IMPORTANT: Base your answer on the specific content provided above. Do not provide generic legal advice - reference the actual text in the document.

Please provide:
1. A direct answer to the question based on THIS document's content
2. Reference to relevant sections of THIS document that support your answer
3. Any important context or implications from THIS document
4. Practical advice based on THIS document's specific terms

Keep your answer accessible to someone without legal training, but make sure it's accurate to this document's actual content."""

        question_message = UserMessage(text=question_prompt)
        
        answer = await send_message(question_prompt)
        if not answer or len(answer.strip()) < 10:
            raise Exception("Empty or invalid response from AI service")
        
        # Save chat message
        chat_message = ChatMessage(
            document_id=request.document_id,
            session_id=session_id,
            question=request.question,
            answer=answer
        )

        chat_dict = chat_message.model_dump()
        chat_dict['timestamp'] = chat_dict['timestamp'].isoformat()

        # Use synchronous MongoDB client to avoid event loop conflicts
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        sync_db.chat_messages.insert_one(chat_dict)
        sync_client.close()
        
        return {
            "question": request.question,
            "answer": answer,
            "session_id": session_id
        }
        
    except Exception as e:
        logging.error(f"Question answering error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to answer question: {str(e)}")

# Get chat history for a document
@api_router.get("/documents/{document_id}/chat")
async def get_chat_history(document_id: str, session_id: Optional[str] = None):
    """Get chat history for a document"""
    try:
        query = {"document_id": document_id}
        if session_id:
            query["session_id"] = session_id

        # Use synchronous MongoDB client to avoid event loop conflicts
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        messages = list(sync_db.chat_messages.find(query).sort("timestamp", 1))
        sync_client.close()

        for msg in messages:
            if isinstance(msg.get('timestamp'), str):
                msg['timestamp'] = datetime.fromisoformat(msg['timestamp'])

        return [ChatMessage(**msg) for msg in messages]

    except Exception as e:
        logging.error(f"Get chat history error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve chat history")

# Export report endpoint
@api_router.post("/documents/{document_id}/export")
async def export_document(document_id: str, export_data: dict):
    """Export document analysis as PDF or Word document"""
    try:
        # Get document from database
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        document = sync_db.documents.find_one({"id": document_id})
        if not document:
            sync_client.close()
            raise HTTPException(status_code=404, detail="Document not found")

        # Check if analysis is completed
        if document.get('analysis_status') != 'completed':
            sync_client.close()
            raise HTTPException(status_code=400, detail="Document analysis not completed")

        # Get export parameters
        format_type = export_data.get('format', 'pdf')
        sections = export_data.get('sections', {})

        # Prepare export content
        export_content = f"""LEGAL DOCUMENT ANALYSIS REPORT

DOCUMENT INFORMATION
===================
Document Name: {document['filename']}
Analysis Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}
Document ID: {document_id}

"""

        # Add sections based on user selection
        if sections.get('summary', True):
            export_content += f"""
EXECUTIVE SUMMARY
=================
{document.get('summary', 'No summary available')}

"""

        if sections.get('clauses', True):
            export_content += f"""
KEY CLAUSES ANALYSIS
====================
"""
            key_clauses = document.get('key_clauses', [])
            for i, clause in enumerate(key_clauses, 1):
                export_content += f"""
{i}. {clause.get('reference', 'N/A')}
   Explanation: {clause.get('explanation', 'N/A')}
   Legal Implications: {clause.get('implications', 'N/A')}
   Business Impact: {clause.get('impact', 'N/A')}
   Potential Concerns: {clause.get('concerns', 'N/A')}

"""

        if sections.get('risks', True):
            export_content += f"""
RISK ASSESSMENT
===============
{document.get('risk_assessment', 'No risk assessment available')}

"""

        if sections.get('qa', True):
            # Get Q&A history
            qa_messages = list(sync_db.chat_messages.find({"document_id": document_id}).sort("timestamp", 1))
            if qa_messages:
                export_content += f"""
Q&A HISTORY
===========
"""
                for msg in qa_messages:
                    export_content += f"""
Question: {msg.get('question', 'N/A')}
Answer: {msg.get('answer', 'N/A')}
Timestamp: {msg.get('timestamp', 'N/A')}

"""

        sync_client.close()

        # Clean the AI response content to remove artifacts
        def clean_ai_response(text):
            if not text:
                return text
            # Remove various AI artifacts and formatting
            text = text.replace('**', '')
            text = text.replace('*', '')
            text = text.replace('```', '')
            text = text.replace('`', '')
            text = text.replace('###', '')
            text = text.replace('##', '')
            text = text.replace('#', '')
            text = text.replace('•', '')
            text = text.replace('·', '')
            text = text.replace('▪', '')
            text = text.replace('◦', '')
            # Remove markdown-style bold and italic
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'_(.*?)_', r'\1', text)
            text = re.sub(r'__(.*?)__', r'\1', text)
            # Remove headers
            text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
            # Clean up excessive whitespace
            text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple newlines to double
            text = re.sub(r'\s+', ' ', text)  # Multiple spaces to single
            text = text.strip()
            return text

        # Apply cleaning to all sections
        summary = clean_ai_response(document.get('summary', ''))
        risk_assessment = clean_ai_response(document.get('risk_assessment', ''))

        # Also clean key clauses
        key_clauses = document.get('key_clauses', [])
        for clause in key_clauses:
            for key in clause:
                if isinstance(clause[key], str):
                    clause[key] = clean_ai_response(clause[key])

        # Generate file content based on format
        if format_type == 'pdf':
            # For PDF, create a formatted document with color coding for risk levels
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.colors import red, orange, green
            from io import BytesIO

            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()

            # Create custom styles for risk levels
            high_risk_style = ParagraphStyle(
                'HighRisk',
                parent=styles['Normal'],
                textColor=red,
                fontSize=12,
                spaceAfter=12
            )
            medium_risk_style = ParagraphStyle(
                'MediumRisk',
                parent=styles['Normal'],
                textColor=orange,
                fontSize=12,
                spaceAfter=12
            )
            low_risk_style = ParagraphStyle(
                'LowRisk',
                parent=styles['Normal'],
                textColor=green,
                fontSize=12,
                spaceAfter=12
            )

            story = []

            # Title
            story.append(Paragraph("LEGAL DOCUMENT ANALYSIS REPORT", styles['Title']))
            story.append(Spacer(1, 24))

            # Document Information
            story.append(Paragraph("DOCUMENT INFORMATION", styles['Heading2']))
            story.append(Paragraph(f"Document Name: {document['filename']}", styles['Normal']))
            story.append(Paragraph(f"Analysis Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}", styles['Normal']))
            story.append(Paragraph(f"Document ID: {document_id}", styles['Normal']))
            story.append(Spacer(1, 12))

            # Add sections
            if sections.get('summary', True):
                story.append(Paragraph("EXECUTIVE SUMMARY", styles['Heading2']))
                story.append(Paragraph(summary, styles['Normal']))
                story.append(Spacer(1, 12))

            if sections.get('clauses', True):
                story.append(Paragraph("KEY CLAUSES ANALYSIS", styles['Heading2']))
                key_clauses = document.get('key_clauses', [])
                for i, clause in enumerate(key_clauses, 1):
                    story.append(Paragraph(f"{i}. {clause.get('reference', 'N/A')}", styles['Heading3']))
                    story.append(Paragraph(f"Explanation: {clause.get('explanation', 'N/A')}", styles['Normal']))
                    story.append(Paragraph(f"Legal Implications: {clause.get('implications', 'N/A')}", styles['Normal']))
                    story.append(Paragraph(f"Business Impact: {clause.get('impact', 'N/A')}", styles['Normal']))
                    story.append(Paragraph(f"Potential Concerns: {clause.get('concerns', 'N/A')}", styles['Normal']))
                    story.append(Spacer(1, 6))

            if sections.get('risks', True):
                story.append(Paragraph("RISK ASSESSMENT", styles['Heading2']))

                # Parse and color-code risk sections
                risk_lines = risk_assessment.split('\n')

                for line in risk_lines:
                    line = line.strip()
                    if not line:
                        continue

                    # Determine color based on risk level keywords
                    if any(keyword in line.upper() for keyword in ['HIGH-RISK', 'CRITICAL', 'SEVERE']):
                        story.append(Paragraph(line, high_risk_style))
                    elif any(keyword in line.upper() for keyword in ['MEDIUM-RISK', 'MODERATE']):
                        story.append(Paragraph(line, medium_risk_style))
                    elif any(keyword in line.upper() for keyword in ['LOW-RISK', 'MINIMAL']):
                        story.append(Paragraph(line, low_risk_style))
                    else:
                        story.append(Paragraph(line, styles['Normal']))

                story.append(Spacer(1, 12))

            if sections.get('qa', True):
                qa_messages = list(sync_db.chat_messages.find({"document_id": document_id}).sort("timestamp", 1))
                if qa_messages:
                    story.append(Paragraph("Q&A HISTORY", styles['Heading2']))
                    for msg in qa_messages:
                        story.append(Paragraph(f"Question: {msg.get('question', 'N/A')}", styles['Normal']))
                        story.append(Paragraph(f"Answer: {msg.get('answer', 'N/A')}", styles['Normal']))
                        story.append(Spacer(1, 6))

            doc.build(story)
            pdf_content = buffer.getvalue()
            buffer.close()

            return Response(
                content=pdf_content,
                media_type='application/pdf',
                headers={"Content-Disposition": f"attachment; filename={document['filename'].replace('.pdf', '')}_analysis.pdf"}
            )

        elif format_type == 'word':
            # For Word document, create a formatted DOCX with color coding
            from docx import Document
            from docx.shared import RGBColor
            from io import BytesIO

            doc = Document()
            doc.add_heading('Legal Document Analysis Report', 0)

            # Add document info
            doc.add_heading('Document Information', level=1)
            doc.add_paragraph(f"Document Name: {document['filename']}")
            doc.add_paragraph(f"Analysis Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
            doc.add_paragraph(f"Document ID: {document_id}")

            # Add sections
            if sections.get('summary', True):
                doc.add_heading('Executive Summary', level=1)
                doc.add_paragraph(summary)

            if sections.get('clauses', True):
                doc.add_heading('Key Clauses Analysis', level=1)
                key_clauses = document.get('key_clauses', [])
                for i, clause in enumerate(key_clauses, 1):
                    doc.add_heading(f'{i}. {clause.get("reference", "N/A")}', level=2)
                    doc.add_paragraph(f"Explanation: {clause.get('explanation', 'N/A')}")
                    doc.add_paragraph(f"Legal Implications: {clause.get('implications', 'N/A')}")
                    doc.add_paragraph(f"Business Impact: {clause.get('impact', 'N/A')}")
                    doc.add_paragraph(f"Potential Concerns: {clause.get('concerns', 'N/A')}")

            if sections.get('risks', True):
                doc.add_heading('Risk Assessment', level=1)

                # Parse and color-code risk sections line by line
                risk_lines = risk_assessment.split('\n')

                for line in risk_lines:
                    line = line.strip()
                    if not line:
                        continue

                    p = doc.add_paragraph()

                    # Determine color based on risk level keywords
                    if any(keyword in line.upper() for keyword in ['HIGH-RISK', 'CRITICAL', 'SEVERE']):
                        run = p.add_run(line)
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                    elif any(keyword in line.upper() for keyword in ['MEDIUM-RISK', 'MODERATE']):
                        run = p.add_run(line)
                        run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
                    elif any(keyword in line.upper() for keyword in ['LOW-RISK', 'MINIMAL']):
                        run = p.add_run(line)
                        run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                    else:
                        p.add_run(line)

            if sections.get('qa', True):
                doc.add_heading('Q&A History', level=1)
                qa_messages = list(sync_db.chat_messages.find({"document_id": document_id}).sort("timestamp", 1))
                for msg in qa_messages:
                    doc.add_paragraph(f"Question: {msg.get('question', 'N/A')}")
                    doc.add_paragraph(f"Answer: {msg.get('answer', 'N/A')}")

            buffer = BytesIO()
            doc.save(buffer)
            docx_content = buffer.getvalue()
            buffer.close()

            return Response(
                content=docx_content,
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={"Content-Disposition": f"attachment; filename={document['filename'].replace('.docx', '').replace('.pdf', '')}_analysis.docx"}
            )

        else:
            raise HTTPException(status_code=400, detail="Unsupported export format")

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Export error: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

# Then include the router in the main app
app.include_router(api_router)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
